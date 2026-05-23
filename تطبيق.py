import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
import pandas as pd
from datetime import datetime
# ==========================================
# CONFIGURATION DE LA PAGE (DOIT ÊTRE EN PREMIER)
# ==========================================
st.set_page_config(
    page_title="Plateforme de gestion des EDTs-S2-2026-Département d'Électrotechnique-Faculté de génie électrique-UDL-SBA",
    layout="wide"  # Garde l'affichage large actuel de votre interface
)

# ==========================================
# INJECTION DU STYLE POUR CACHER STREAMLIT
# ==========================================
hide_streamlit_style = """
    <style>
    /* Masquer le menu hamburger, le bouton de déploiement et la barre supérieure */
    header {visibility: hidden;}
    [data-testid="stToolbar"] {visibility: hidden; height: 0px;}
    [data-testid="stDecoration"] {visibility: hidden; height: 0px;}
    [data-testid="stStatusWidget"] {visibility: hidden; height: 0px;}
    #MainMenu {visibility: hidden;}
    
    /* Masquer le pied de page "Made with Streamlit" */
    footer {visibility: hidden;}
    
    /* Réduire l'espace blanc du haut créé par la disparition du header */
    .block-container {
        padding-top: 2rem;
    }
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
# ==========================================
# CONFIGURATION ET CONSTANTES
# ==========================================
TITRE_PLATEFORME = "Plateforme de gestion des EDTs & assiduité des étudiants_2025-2026-Département d'Électrotechnique-Faculté de génie électrique-UDL-SBA"

DEPARTEMENTS = [
    "Département d'Électrotechnique",
    "Département d'Électronique",
    "Département d'Automatique",
    "Département de Télécommunications"
]

TYPES_DOCUMENTS = [
    "Bordereau d'envoi",
    "Justification d'absence",
    "Procès-verbal (PV) de réunion",
    "PV de surveillance",
    "PV du Comité Pédagogique",
    "Réunion du Conseil de Discipline"
]

OPTIONS_DESTINATAIRES = [
    "Le Doyen de la faculté",
    "Le vice Doyen de la Post graduation",
    "Le vice Doyen de la graduation",
    "Autres"
]

MOTIFS_ABSENCE = [
    "Personnel",
    "Médical",
    "Autre"
]

LISTE_SPECIALITES = [
    "Réseaux électriques",
    "Energies renouvelables",
    "Commandes électriques",
    "Master-MCIL",
    "Licence_MCIL",
    "Licence_ELT",
    "Ingénieur",
    "Ingénieur_EI",
    "Ingénieur_RSE"
]

LISTE_ANNEES_ETUDE = [
    "1ère Année Master",
    "2ème Année Master",
    "1ère Année",
    "2ème Année",
    "3ème Année",
    "4ème Année",
    "5ème Année"
]

NOM_FICHIER_EXCEL = "dataEDT-ELT-S2-2026.xlsx"

COLONNES_SUIVI_OFFICIELRES = [
    "Nom & Prénom Étudiant",
    "Date de l'absence",
    "Date de délivrance",
    "Enseignant",
    "Matière",
    "Nombre d'absences comptabilisées",
    "Situation Réglementaire"
]

# ==========================================
# CHARGEMENT ET TRAITEMENT DU FICHIER SOURCE EXCEL
# ==========================================
@st.cache_data
def charger_base_enseignements(chemin_fichier):
    """
    Charge le fichier Excel source et structure un dictionnaire indexé par Enseignant.
    La disposition originale est respectée : Enseignements, Code, Enseignants, Horaire, Jours, Lieu, Promotion.
    """
    dict_enseignants = {}
    if os.path.exists(chemin_fichier):
        try:
            df = pd.read_excel(chemin_fichier)
            df.columns = [col.strip() for col in df.columns]
            
            if "Enseignants" in df.columns and "Enseignements" in df.columns and "Code" in df.columns:
                for _, row in df.iterrows():
                    nom_ens = str(row["Enseignants"]).strip()
                    matiere_nom = str(row["Enseignements"]).strip()
                    code_matiere = str(row["Code"]).strip()
                    
                    if nom_ens and nom_ens != "nan" and matiere_nom and matiere_nom != "nan":
                        libelle_matiere = f"{matiere_nom} ({code_matiere})"
                        
                        if nom_ens not in dict_enseignants:
                            dict_enseignants[nom_ens] = []
                        if libelle_matiere not in dict_enseignants[nom_ens]:
                            dict_enseignants[nom_ens].append(libelle_matiere)
            else:
                st.error("Format critique absent : Les colonnes 'Enseignants', 'Enseignements' ou 'Code' manquent dans le fichier Excel.")
        except Exception as e:
            st.error(f"Erreur lors de la lecture du fichier Excel source : {str(e)}")
            
    if not dict_enseignants:
        dict_enseignants = {
            "Zidi": ["Stabilité et dynamique des réseaux électriques (Cours-SDRE-RE)"],
            "Bermaki": ["Éclairage LED: Principes et applications (Cours-LEDPA-RE)"],
            "Touhami": ["Techniques d'intelligence artificielle (Cours-TIA-RE)"],
            "BENHAMIDA": ["Intégration des ressources renouvelables aux réseaux électriques (Cours-IRRRE-RE)"],
            "Rezoug": ["Dimensionnement des Réseaux électriques industriels (Cours-DREI-RE)"],
            "Bellebna": ["Technique de la haute tension (Cours-THT-RE)"],
            "Benhamida": ["Conduite des réseaux électriques (Cours-CdRE-RE)"],
            "Maamar": ["Réseaux électriques intelligents (Cours-REI-RE)"]
        }
    return dict_enseignants

DATA_ENSEIGNANTS = charger_base_enseignements(NOM_FICHIER_EXCEL)

# ==========================================
# INITIALISATION DU SUIVI ET HISTORIQUE (SESSION STATE)
# ==========================================
if "historique_justifications" not in st.session_state:
    st.session_state["historique_justifications"] = []

if "compteur_absences_strict" not in st.session_state:
    st.session_state["compteur_absences_strict"] = {}

if "metadonnees_absences_strict" not in st.session_state:
    st.session_state["metadonnees_absences_strict"] = {}

# ==========================================
# FONCTIONS TECHNIQUES DE STRUCTURE DOCUMENTAIRE
# ==========================================
def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    """Définit l'espacement interne (padding) compact des cellules d'un tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def appliquer_bordure_cellule_noire(cell):
    """Applique une bordure fine noire standard et complète sur tous les côtés d'une cellule."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def ajouter_champ_page(run, type_champ):
    """Injecte un champ de numérotation dynamique (PAGE ou NUMPAGES) dans un paragraphe Word."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = type_champ
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def initialiser_paragraphe_strict(p):
    """Supprime totalement les espaces et configure un interligne simple strict."""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

def appliquer_structure_pages_sans_ref(doc):
    """Configure les marges globales et positionne le numéro de page à l'extrémité droite du pied de page."""
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        section.different_first_page_header_footer = False
        
        footer = section.footer
        footer_p = footer.paragraphs[0]
        initialiser_paragraphe_strict(footer_p)
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        r_page_actuelle = footer_p.add_run()
        r_page_actuelle.font.name = 'Calibri'
        r_page_actuelle.font.size = Pt(11)
        ajouter_champ_page(r_page_actuelle, "PAGE")
        
        r_separateur = footer_p.add_run(" / ")
        r_separateur.font.name = 'Calibri'
        r_separateur.font.size = Pt(12)
        
        r_total_pages = footer_p.add_run()
        r_total_pages.font.name = 'Calibri'
        r_total_pages.font.size = Pt(11)
        ajouter_champ_page(r_total_pages, "NUMPAGES")

def inserer_bloc_en_tete_bordereau(doc, departement):
    """Génère l'en-tête standard classique avec logo pour le Bordereau d'envoi."""
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.19)
    header_table.columns[1].width = Inches(5.1)
    
    cell_logo = header_table.rows[0].cells[0]
    cell_texte = header_table.rows[0].cells[1]
    
    tblPr = header_table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    p_logo = cell_logo.paragraphs[0]
    initialiser_paragraphe_strict(p_logo)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    nom_fichier_logo = "logo.PNG"
    if os.path.exists(nom_fichier_logo):
        p_logo.add_run().add_picture(nom_fichier_logo, width=Inches(0.58), height=Inches(0.94))
    else:
        r_alt = p_logo.add_run("[LOGO]")
        r_alt.font.name = 'Calibri'
        r_alt.font.size = Pt(11)
        r_alt.font.italic = True

    p_en_tete = cell_texte.paragraphs[0]
    initialiser_paragraphe_strict(p_en_tete)
    p_en_tete.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = p_en_tete.add_run("RÉPUBLIQUE ALGÉRIENNE DÉMOCRATIQUE ET POPULAIRE\n")
    r1.bold = True
    
    r2 = p_en_tete.add_run(
        "Ministère de l'Enseignement Supérieur et de la Recherche Scientifique\n"
        "Université Djillali Liabes - Sidi Bel Abbès\n"
        "Faculté de Génie Électrique\n"
    )
    
    r_dept = p_en_tete.add_run(f"{departement.upper()}")
    r_dept.bold = True
    
    for r in [r1, r2, r_dept]:
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
    
    p_espace = doc.add_paragraph()
    initialiser_paragraphe_strict(p_espace)

# ==========================================
# GÉNÉRATEUR : JUSTIFICATION D'ABSENCE
# ==========================================
def generer_justificatif_iso(departement, donnees):
    doc = Document()
    appliquer_structure_pages_sans_ref(doc)
    
    grid_table = doc.add_table(rows=2, cols=3)
    grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid_table.autofit = False
    grid_table.allow_autofit = False
    
    dxa_widths = [int(1.19 * 1440), int(3.70 * 1440), int(1.40 * 1440)]
    inch_widths = [Inches(1.19), Inches(3.70), Inches(1.40)]
    
    tblPr = grid_table._tbl.tblPr
    tblGrid = OxmlElement('w:tblGrid')
    for width_dxa in dxa_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width_dxa))
        tblGrid.append(gridCol)
    grid_table._tbl.insert(1, tblGrid)

    for row in grid_table.rows:
        for idx, width_inch in enumerate(inch_widths):
            cell = row.cells[idx]
            cell.width = width_inch
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(dxa_widths[idx]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    cell_logo_haut = grid_table.cell(0, 0)
    cell_logo_bas = grid_table.cell(1, 0)
    cell_univ_haut = grid_table.cell(0, 1)
    cell_titre_bas = grid_table.cell(1, 1)
    cell_meta_haut = grid_table.cell(0, 2)
    cell_meta_bas = grid_table.cell(1, 2)
    
    for row in grid_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            appliquer_bordure_cellule_noire(cell)
            for p in cell.paragraphs:
                initialiser_paragraphe_strict(p)

    cell_logo_haut.merge(cell_logo_bas)
    cell_meta_haut.merge(cell_meta_bas)
    
    appliquer_bordure_cellule_noire(cell_logo_haut)
    appliquer_bordure_cellule_noire(cell_logo_bas)
    appliquer_bordure_cellule_noire(cell_meta_haut)
    appliquer_bordure_cellule_noire(cell_meta_bas)

    p_logo = cell_logo_haut.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nom_fichier_logo = "logo.PNG"
    if os.path.exists(nom_fichier_logo):
        p_logo.add_run().add_picture(nom_fichier_logo, width=Inches(0.58), height=Inches(0.94))
    else:
        r_alt = p_logo.add_run("[ LOGO ]")
        r_alt.font.name = 'Calibri'
        r_alt.font.size = Pt(11)
        r_alt.font.italic = True

    p_univ = cell_univ_haut.paragraphs[0]
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re1 = p_univ.add_run("Université Djillali Liabes\n")
    re1.bold = True
    re2 = p_univ.add_run("Sidi Bel Abbes")
    for re in [re1, re2]:
        re.font.name = 'Calibri'
        re.font.size = Pt(11)

    p_titre = cell_titre_bas.paragraphs[0]
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_titre.add_run("JUSTIFICATION D’ABSENCE")
    rc.font.name = 'Arial'
    rc.font.size = Pt(11)
    rc.bold = True

    p_meta = cell_meta_haut.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rm1 = p_meta.add_run("Code : PPER.06\n")
    rm2 = p_meta.add_run("Révision : 00\n")
    rm3 = p_meta.add_run("Date : 16/05/2026\n")
    rm4 = p_meta.add_run("Pages : 1/1")
    for rm in [rm1, rm2, rm3, rm4]:
        rm.font.name = 'Calibri'
        rm.font.size = Pt(11)

    p_esp1 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp1)
    
    p_fac = doc.add_paragraph()
    initialiser_paragraphe_strict(p_fac)
    rfac = p_fac.add_run("Faculté de génie Electrique\n")
    rfac.bold = True
    rdept_txt = p_fac.add_run(f"{departement}")
    for rf in [rfac, rdept_txt]:
        rf.font.name = 'Calibri'
        rf.font.size = Pt(11)
        
    p_esp2 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp2)
    
    p_corps = doc.add_paragraph()
    initialiser_paragraphe_strict(p_corps)
    p_corps.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_corps = p_corps.add_run(f"Le {departement} atteste par la présente que l'étudiant(e) :")
    r_corps.font.name = 'Calibri'
    r_corps.font.size = Pt(11)
    
    p_esp3 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp3)
    
    p_id = doc.add_paragraph()
    initialiser_paragraphe_strict(p_id)
    
    r_nom_lbl = p_id.add_run("Nom et prénom : ")
    r_nom_lbl.bold = True
    r_nom_val = p_id.add_run(f"{donnees['nom_prenom']}\n")
    
    r_annee_lbl = p_id.add_run("Année d’étude : ")
    r_annee_lbl.bold = True
    r_annee_val = p_id.add_run(f"{donnees['annee_etude']}\n")
    
    r_spec_lbl = p_id.add_run("Spécialité : ")
    r_spec_lbl.bold = True
    r_spec_val = p_id.add_run(f"{donnees['specialite']}\n")
    
    r_ens_lbl = p_id.add_run("Enseignant concerné : ")
    r_ens_lbl.bold = True
    r_ens_val = p_id.add_run(f"Pr. {donnees['enseignant']}\n")
    
    r_mat_lbl = p_id.add_run("Matière / Enseignement : ")
    r_mat_lbl.bold = True
    r_mat_val = p_id.add_run(f"{donnees['matiere']}\n")
    
    r_abs_lbl = p_id.add_run("a été absent(e) durant la période allant du : ")
    r_abs_lbl.bold = True
    date_deb_txt = donnees['date_debut'].strftime('%d/%m/%Y')
    date_fin_txt = donnees['date_fin'].strftime('%d/%m/%Y')
    r_abs_val = p_id.add_run(f"{date_deb_txt} au {date_fin_txt}")
    
    for run in [r_nom_lbl, r_nom_val, r_annee_lbl, r_annee_val, r_spec_lbl, r_spec_val, r_ens_lbl, r_ens_val, r_mat_lbl, r_mat_val, r_abs_lbl, r_abs_val]:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        
    p_esp4 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp4)
    
    p_motif_titre = doc.add_paragraph()
    initialiser_paragraphe_strict(p_motif_titre)
    r_mot_titre = p_motif_titre.add_run("Pour le motif suivant :")
    r_mot_titre.bold = True
    r_mot_titre.font.name = 'Calibri'
    r_mot_titre.font.size = Pt(12)
    
    for motif in MOTIFS_ABSENCE:
        p_m = doc.add_paragraph()
        initialiser_paragraphe_strict(p_m)
        p_m.paragraph_format.left_indent = Inches(0.4)
        if motif == donnees['motif_selectionne']:
            r_box = p_m.add_run("[ X ]  ")
            r_box.bold = True
        else:
            r_box = p_m.add_run("[   ]  ")
        r_txt = p_m.add_run(motif)
        
        r_box.font.name = 'Calibri'
        r_box.font.size = Pt(12)
        r_txt.font.name = 'Calibri'
        r_txt.font.size = Pt(12)
        
    p_esp5 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp5)
    
    p_cloture = doc.add_paragraph()
    initialiser_paragraphe_strict(p_cloture)
    r_cloture = p_cloture.add_run("La présente attestation est délivrée à l’intéressé(e) pour servir et valoir ce que de droit.")
    r_cloture.font.name = 'Calibri'
    r_cloture.font.size = Pt(11)
    r_cloture.font.italic = True
    
    p_esp6 = doc.add_paragraph()
    initialiser_paragraphe_strict(p_esp6)
    
    p_sig = doc.add_paragraph()
    initialiser_paragraphe_strict(p_sig)
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_edit_txt = donnees['date_edition'].strftime('%d/%m/%Y')
    run_sig = p_sig.add_run(f"Fait à : Sidi Bel Abbès.\t\tLe : {date_edit_txt}\n\n\t\t\t\t\t\tLe Chef de Département")
    run_sig.font.name = 'Calibri'
    run_sig.font.size = Pt(11)
    run_sig.bold = True
    
    return doc

# ==========================================
# GÉNÉRATEUR : BORDEREAU D'ENVOI
# ==========================================
def generer_bordereau_iso(departement, donnees):
    doc = Document()
    
    # Configuration des marges globales de la page (0.8 pouce partout)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Propagation du pied de page sur toutes les pages
        section.different_first_page_header_footer = False
        
        # Structure du pied de page rectifié
        footer = section.footer
        footer_p = footer.paragraphs[0]
        
        footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_pPr = footer_p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        
        # 1. Taquet au centre pour la référence (Centre à ~ 3.45 pouces = 4968 dxa)
        tab_centre = OxmlElement('w:tab')
        tab_centre.set(qn('w:val'), 'center')
        tab_centre.set(qn('w:pos'), '4968')
        tabs.append(tab_centre)
        
        # 2. Taquet à l'extrême droite pour les numéros de page (Extrémité à 6.9 pouces = 9936 dxa)
        tab_droite = OxmlElement('w:tab')
        tab_droite.set(qn('w:val'), 'right')
        tab_droite.set(qn('w:pos'), '9936')
        tabs.append(tab_droite)
        
        footer_pPr.append(tabs)
        
        # Premier saut vers le centre pour y écrire le code de référence
        footer_p.add_run("\t")
        r_ref_fixe = footer_p.add_run("Réf : UDL-GEL-ER-004-2026")
        r_ref_fixe.font.name = 'Calibri'
        r_ref_fixe.font.size = Pt(11)
        
        # Deuxième saut vers l'extrême droite pour y loger la pagination automatique
        footer_p.add_run("\t")
        
        r_page_actuelle = footer_p.add_run()
        r_page_actuelle.font.name = 'Calibri'
        r_page_actuelle.font.size = Pt(11)
        ajouter_champ_page(r_page_actuelle, "PAGE")
        
        r_separateur = footer_p.add_run("/")
        r_separateur.font.name = 'Calibri'
        r_separateur.font.size = Pt(11)
        
        r_total_pages = footer_p.add_run()
        r_total_pages.font.name = 'Calibri'
        r_total_pages.font.size = Pt(11)
        ajouter_champ_page(r_total_pages, "NUMPAGES")

    # 1. STRUCTURE DE L'EN-TÊTE VIA UN TABLEAU INVISIBLE
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.autofit = False
    
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(5.7)
    
    cell_logo = header_table.rows[0].cells[0]
    cell_texte = header_table.rows[0].cells[1]
    
    tblPr = header_table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Insertion du Logo (Largeur 80 pixels = 0.833 pouces)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    nom_fichier_logo = "logo.PNG"
    if os.path.exists(nom_fichier_logo):
        p_logo.add_run().add_picture(nom_fichier_logo, width=Inches(0.833))
    else:
        r_alt = p_logo.add_run("[LOGO UNIVERSITÉ]")
        r_alt.font.name = 'Calibri'
        r_alt.font.size = Pt(8)
        r_alt.font.italic = True

    # Insertion des textes officiels de l'en-tête (Calibri)
    p_en_tete = cell_texte.paragraphs[0]
    p_en_tete.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = p_en_tete.add_run("RÉPUBLIQUE ALGÉRIENNE DÉMOCRATIQUE ET POPULAIRE\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    
    r2 = p_en_tete.add_run(
        "Ministère de l'Enseignement Supérieur et de la Recherche Scientifique\n"
        "Université Djillali Liabes - Sidi Bel Abbès\n"
        "Faculté de Génie Électrique\n"
    )
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    
    r_dept = p_en_tete.add_run(f"{departement.upper()}\n")
    r_dept.bold = True
    r_dept.font.size = Pt(11)
    r_dept.font.name = 'Calibri'

    doc.add_paragraph("\n")

    # 2. RÉFÉRENCE CHRONOLOGIQUE
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_ref = p_ref.add_run(f"N° : {donnees['num_reference']}/ F.G.E/ V.D.E.Q.L.E/2026")
    r_ref.font.size = Pt(10)
    r_ref.font.name = 'Calibri'
    r_ref.bold = True

    doc.add_paragraph("\n")

    # 3. TITRE DU BORDEREAU (Taille 36, Calibri, Italique, Souligné)
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titre = p_titre.add_run("BORDEREAU D’ENVOI")
    r_titre.font.name = 'Calibri'
    r_titre.font.size = Pt(36)
    r_titre.italic = True
    r_titre.underline = True
    r_titre.bold = True
    
    doc.add_paragraph("\n")

    # 4. DESTINATAIRE CONSTRUIT DYNAMIQUEMENT (Calibri)
    p_dest = doc.add_paragraph()
    p_dest.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_dest = p_dest.add_run(f"A monsieur : {donnees['destinataire']}")
    r_dest.bold = True
    r_dest.font.size = Pt(12)
    r_dest.font.name = 'Calibri'

    doc.add_paragraph("\n")

    # 5. TABLEAU DE TRANSMISSION MULTI-LIGNES
    liste_pieces = donnees['liste_pieces']
    nb_lignes_totatles = 2 + len(liste_pieces)
    
    table = doc.add_table(rows=nb_lignes_totatles, cols=3)
    table.style = 'Table Grid'
    
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(0.8)
    table.columns[2].width = Inches(1.7)

    # Ligne 1 : En-têtes fixes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Désignation des pièces"
    hdr_cells[1].text = "Nbre"
    hdr_cells[2].text = "Observations"
    
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Calibri'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_margins(cell, top=120, bottom=120)

    # Ligne 2 : Formule d'accompagnement
    row_joint = table.rows[1].cells
    row_joint[0].text = "Veuillez trouver ci-joint :"
    row_joint[0].paragraphs[0].runs[0].font.italic = True
    row_joint[0].paragraphs[0].runs[0].font.name = 'Calibri'
    row_joint[0].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_margins(row_joint[0], top=80, bottom=80)

    # Lignes Dynamiques
    for index, piece in enumerate(liste_pieces):
        row_idx = 2 + index
        current_row = table.rows[row_idx].cells
        
        current_row[0].text = str(piece["Désignation des pièces"])
        current_row[1].text = str(piece["Nbre"])
        current_row[2].text = str(piece["Observations"])
        
        for i, cell in enumerate(current_row):
            set_cell_margins(cell, top=150, bottom=300)
            if len(cell.paragraphs[0].runs) > 0:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            if i == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n\n")

    # 6. SIGNATURES ET ACCUSÉ DE RÉCEPTION
    p_signatures = doc.add_paragraph()
    p_signatures.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_texte = donnees['date_creation'].strftime('%d/%m/%Y')
    run_sig = p_signatures.add_run(f"Sidi bel Abbès le : {date_texte}\t\t\t\tChef de département")
    run_sig.font.name = 'Calibri'
    run_sig.font.size = Pt(11)
    run_sig.bold = True

    doc.add_paragraph("\n\n\n\n")

    p_accuse = doc.add_paragraph()
    p_accuse.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_accuse = p_accuse.add_run("Accusé de réception    ")
    run_accuse.font.name = 'Calibri'
    run_accuse.font.size = Pt(10)
    run_accuse.font.underline = True
    run_accuse.bold = True

    return doc

def generer_pv_generique(departement, type_pv, donnees):
    """Générateur secondaire de secours (Calibri)."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(f"{type_pv} - {departement}\nDocument en cours.")
    run.font.name = 'Calibri'
    return doc

# ==========================================
# INTERFACE UTILISATEUR STREAMLIT
# ==========================================
st.set_page_config(page_title="Générateur ISO Multi-Documents", layout="wide")

st.caption(TITRE_PLATEFORME)
st.title("Gestion Administrative - Générateurs de documents")

col_dept, col_doc = st.columns(2)
with col_dept:
    dept_choisi = st.selectbox("Département émetteur :", DEPARTEMENTS)
with col_doc:
    doc_choisi = st.selectbox("Nature du document à générer :", TYPES_DOCUMENTS, index=1)

st.divider()

tab_formulaire, tab_historique = st.tabs(["📝 Édition du Document", "📊 Historique & Compteur d'Exclusions"])

with tab_formulaire:
    st.subheader(f"Formulaire d'édition - {doc_choisi}")
    donnees_doc = {}
    verrou_dates_invalides = False

    # --- FORMULAIRE : BORDEREAU D'ENVOI ---
    if doc_choisi == "Bordereau d'envoi":
        col_ref, col_date = st.columns(2)
        with col_ref:
            donnees_doc['num_reference'] = st.text_input("Référence séquentielle (Ex: 27)", value="27")
        with col_date:
            donnees_doc['date_creation'] = st.date_input("Date d'édition", datetime.now())
            
        st.markdown("##### Destinataire officiel")
        choix_dest = st.selectbox("Sélectionnez le destinataire dans la liste :", OPTIONS_DESTINATAIRES, index=0)
        
        if choix_dest == "Autres":
            donnees_doc['destinataire'] = st.text_input("Veuillez saisir la destination personnalisée :", value="")
        else:
            donnees_doc['destinataire'] = choix_dest
            
        st.markdown("---")
        st.write("**Configuration du Tableau de Transmission**")
        
        df_initial = pd.DataFrame([
            {"Désignation des pièces": "Fiches de vœux du second semestre", "Nbre": 12, "Observations": "Pour examen"},
            {"Désignation des pièces": "Procès-verbal de délibération", "Nbre": 2, "Observations": "Pour affichage"}
        ])
        
        df_edite = st.data_editor(
            df_initial, 
            num_rows="dynamic", 
            use_container_width=True,
            column_config={
                "Désignation des pièces": st.column_config.TextColumn(width="medium", required=True),
                "Nbre": st.column_config.NumberColumn(width="small", min_value=1, required=True),
                "Observations": st.column_config.TextColumn(width="medium")
            }
        )
        donnees_doc['liste_pieces'] = df_edite.to_dict(orient="records")

    import os
    import urllib.parse
    import datetime
    import pandas as pd
    import streamlit as st
    
    # --- ACCÈS AUX ÉTUDIANTS DEPUIS LE FICHIER SOURCE ---
    liste_noms_etudiants = []
    nom_fichier_etudiants = "Liste des étudiants_2025-2026.xlsx"
    github_url_racine = "https://raw.githubusercontent.com/votre_compte/votre_depot/main/"
    
    # Encodage de l'adresse brute pour GitHub (sécurisation espaces et accents)
    nom_fichier_encode = urllib.parse.quote(nom_fichier_etudiants)
    url_complete_etudiants = f"{github_url_racine}{nom_fichier_encode}"
    
    # Tentative de récupération des étudiants (Distant GitHub puis Repli Local)
    df_etudiants_source = None
    try:
        df_etudiants_source = pd.read_excel(url_complete_etudiants)
    except Exception:
        if os.path.exists(nom_fichier_etudiants):
            try:
                df_etudiants_source = pd.read_excel(nom_fichier_etudiants)
            except Exception:
                pass
    
    # Extraction et nettoyage alphabétique de la colonne des noms d'étudiants
    if df_etudiants_source is not None:
        df_etudiants_source.columns = [str(col).strip() for col in df_etudiants_source.columns]
        for col_detectee in ["Nom et prénom", "Nom et Prénom", "Etudiants", "Etudiant", "Nom"]:
            if col_detectee in df_etudiants_source.columns:
                liste_noms_etudiants = df_etudiants_source[col_detectee].dropna().astype(str).str.strip().unique().tolist()
                liste_noms_etudiants.sort()
                break
    
    # Initialisation de la variable de verrouillage de sécurité
    verrou_dates_invalides = False
    
    # --- FORMULAIRE : JUSTIFICATION D'ABSENCE ---
    if doc_choisi == "Justification d'absence":
        col_nom, col_annee = st.columns(2)
        
        with col_nom:
            # Remplacement du text_input par le selectbox alimenté par la liste des étudiants du fichier source
            if liste_noms_etudiants:
                donnees_doc['nom_prenom'] = st.selectbox(
                    "Nom et prénom de l'étudiant(e) :",
                    options=liste_noms_etudiants
                )
            else:
                # Saisie manuelle de secours préservée si le fichier source est inaccessible
                donnees_doc['nom_prenom'] = st.text_input(
                    "Nom et prénom de l'étudiant(e) :", 
                    value="Benali Mohamed"
                )
                st.warning("⚠️ Base de données des étudiants introuvable. Saisie libre activée.")
                
        with col_annee:
            donnees_doc['annee_etude'] = st.selectbox(
                "Année d'étude (Ex: 1ère Année Master) :", 
                LISTE_ANNEES_ETUDE, 
                index=1
            )
            
        col_spec, col_motif = st.columns(2)
        with col_spec:
            donnees_doc['specialite'] = st.selectbox(
                "Spécialité / Option :", 
                LISTE_SPECIALITES, 
                index=0
            )
        with col_motif:
            donnees_doc['motif_selectionne'] = st.selectbox(
                "Motif réglementaire retenu :", 
                MOTIFS_ABSENCE, 
                index=0
            )
            
        st.markdown("##### Informations Complémentaires (Enseignant & Enseignement)")
        col_ens, col_mat = st.columns(2)
        
        with col_ens:
            liste_nom_enseignants = sorted(list(DATA_ENSEIGNANTS.keys()))
            index_par_defaut = liste_nom_enseignants.index("Bellebna") if "Bellebna" in liste_nom_enseignants else 0
            enseignant_choisi = st.selectbox(
                "Enseignant ayant déclaré l'absence :", 
                liste_nom_enseignants, 
                index=index_par_defaut
            )
            donnees_doc['enseignant'] = enseignant_choisi
            
        with col_mat:
            liste_matieres_disponibles = sorted(DATA_ENSEIGNANTS[enseignant_choisi])
            matiere_choisie = st.selectbox(
                "Matière concernée :", 
                liste_matieres_disponibles
            )
            donnees_doc['matiere'] = matiere_choisie
    
        col_d1, col_d2, col_d3 = st.columns(3)
        with col_d1:
            donnees_doc['date_debut'] = st.date_input(
                "Date de début de l'absence", 
                datetime.date.today()
            )
        with col_d2:
            donnees_doc['date_fin'] = st.date_input(
                "Date de fin de l'absence", 
                datetime.date.today()
            )
        with col_d3:
            donnees_doc['date_edition'] = st.date_input(
                "Date de délivrance (Justification)", 
                datetime.date.today()
            )
    
        # Vérification et Verrouillage Strict Chronologique des dates
        if donnees_doc['date_edition'] < donnees_doc['date_debut']:
            verrou_dates_invalides = True
            st.error("🚨 ERREUR CRITIQUE DE SAISIE : La date de délivrance de la justification ne peut pas être antérieure à la date de début de l'absence. Le système bloque la génération du document.")
        
        nom_etudiant_clean = donnees_doc['nom_prenom'].strip()
        if nom_etudiant_clean and not verrou_dates_invalides:
            cle_composite_verif = (nom_etudiant_clean, donnees_doc['enseignant'], donnees_doc['matiere'])
            
            # Sécurisation de l'accès au dictionnaire de session_state
            if "compteur_absences_strict" not in st.session_state:
                st.session_state["compteur_absences_strict"] = {}
                
            nb_absences_actuelles = st.session_state["compteur_absences_strict"].get(cle_composite_verif, 0)
            
            if nb_absences_actuelles >= 5:
                st.error(f"⚠️ ATTENTION CRITIQUE : L'étudiant '{nom_etudiant_clean}' compte déjà {nb_absences_actuelles} absences dans cette matière spécifique ({donnees_doc['matiere']}) avec Pr. {donnees_doc['enseignant']}. Il est déclaré EXCLU de ce cours.")
            elif nb_absences_actuelles > 0:
                st.warning(f"Note : Cet étudiant possède actuellement {nb_absences_actuelles} absence(s) enregistrée(s) pour ce cours précis.")
    
    else:
        with st.form("form_autres"):
            donnees_doc['date_creation'] = st.date_input("Date", datetime.date.today())
            donnees_doc['contenu'] = st.text_area("Contenu textuel")
            st.form_submit_button("Valider la saisie")
    
    st.markdown("<br>", unsafe_allow_html=True)

    # ==========================================
    # GESTION DES ACTIONS FINALES & COMPILATION
    # ==========================================
    if doc_choisi == "Bordereau d'envoi":
        if st.button("Compiler et Générer le Bordereau Officiel"):
            if not donnees_doc['destinataire'].strip():
                st.error("Erreur : Le champ de destination personnalisée ne peut pas être vide.")
            else:
                try:
                    document_final = generer_bordereau_iso(dept_choisi, donnees_doc)
                    output_stream = io.BytesIO()
                    document_final.save(output_stream)
                    output_stream.seek(0)
                    
                    st.success("✓ Bordereau d'envoi généré avec succès.")
                    st.download_button(
                        label="⬇️ Télécharger le Bordereau d'envoi (.docx)",
                        data=output_stream,
                        file_name=f"Bordereau_{dept_choisi.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as error:
                    st.error(f"Échec de l'opération de génération : {str(error)}")

    elif doc_choisi == "Justification d'absence":
        # Le bouton de soumission est totalement inhibé ou rendu inactif si les dates sont incohérentes
        if verrou_dates_invalides:
            st.button("Compiler et Générer la Justification d'Absence", disabled=True, help="Saisie de date erronée.")
        else:
            if st.button("Compiler et Générer la Justification d'Absence", disabled=False):
                if not donnees_doc['nom_prenom'].strip():
                    st.error("Erreur : Le nom de l'étudiant ne peut pas être vide.")
                else:
                    try:
                        nom_etudiant = donnees_doc['nom_prenom'].strip()
                        ens_concerne = donnees_doc['enseignant']
                        mat_concerne = donnees_doc['matiere']
                        
                        cle_composite = (nom_etudiant, ens_concerne, mat_concerne)
                        
                        document_final = generer_justificatif_iso(dept_choisi, donnees_doc)
                        output_stream = io.BytesIO()
                        document_final.save(output_stream)
                        output_stream.seek(0)
                        
                        st.session_state["compteur_absences_strict"][cle_composite] = st.session_state["compteur_absences_strict"].get(cle_composite, 0) + 1
                        total_absences = st.session_state["compteur_absences_strict"][cle_composite]
                        
                        st.session_state["metadonnees_absences_strict"][cle_composite] = {
                            "date_absence": donnees_doc['date_debut'].strftime('%d/%m/%Y'),
                            "date_delivrance": donnees_doc['date_edition'].strftime('%d/%m/%Y'),
                            "enseignant": ens_concerne,
                            "matiere": mat_concerne
                        }
                        
                        enregistrement_historique = {
                            "Date Opération": datetime.now().strftime("%d/%m/%Y %H:%M"),
                            "Étudiant": nom_etudiant,
                            "Année d'étude": donnees_doc['annee_etude'],
                            "Spécialité": donnees_doc['specialite'],
                            "Enseignant": ens_concerne,
                            "Matière": mat_concerne,
                            "Motif": donnees_doc['motif_selectionne'],
                            "Absences dans cette matière": total_absences,
                            "Statut": "EXCLU (Matière)" if total_absences >= 5 else "Actif"
                        }
                        st.session_state["historique_justifications"].append(enregistrement_historique)
                        
                        st.success(f"✓ Justification validée (Absence n°{total_absences} enregistrée pour le cours de {mat_concerne}).")
                        
                        if total_absences >= 5:
                            st.error(f"🚨 Seuil critique atteint ! L'étudiant {nom_etudiant} est déclaré EXCLU du module : {mat_concerne}.")
                            
                        st.download_button(
                            label="⬇️ Télécharger le justificatif (.docx)",
                            data=output_stream,
                            file_name=f"Justification_Absence_{nom_etudiant.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as error:
                        st.error(f"Échec de l'opération de génération : {str(error)}")

# ==========================================
# ONGLET : HISTORIQUE ET COMPTEURS D'EXCLUSIONS
# ==========================================
with tab_historique:
    st.subheader("📋 Historique Global des Étudiants Justifiés")
    
    if len(st.session_state["historique_justifications"]) == 0:
        st.info("Aucune justification d'absence n'a encore été enregistrée lors de cette session.")
    else:
        df_historique = pd.DataFrame(st.session_state["historique_justifications"])
        st.dataframe(df_historique, use_container_width=True)
        
    st.divider()
    st.subheader("🚨 Suivi par Matière et Seuils d'Exclusion (>= 5 Absences par matière)")
    
    donnees_compteur = []
    donnees_exclus_uniquement = []
    
    for triplet_cle, nb_abs in st.session_state["compteur_absences_strict"].items():
        etudiant_nom, ens_nom, mat_nom = triplet_cle
        meta = st.session_state["metadonnees_absences_strict"].get(triplet_cle, {"date_absence": "N/A", "date_delivrance": "N/A", "enseignant": ens_nom, "matiere": mat_nom})
        
        situation = "❌ EXCLU DE LA MATIÈRE" if nb_abs >= 5 else "✅ En règle"
        
        ligne_complete = {
            "Nom & Prénom Étudiant": etudiant_nom,
            "Date de l'absence": meta["date_absence"],
            "Date de délivrance": meta["date_delivrance"],
            "Enseignant": meta["enseignant"],
            "Matière": meta["matiere"],
            "Nombre d'absences comptabilisées": nb_abs,
            "Situation Réglementaire": situation
        }
        donnees_compteur.append(ligne_complete)
        
        if nb_abs >= 5:
            donnees_exclus_uniquement.append(ligne_complete)
        
    if donnees_compteur:
        df_compteur = pd.DataFrame(donnees_compteur)
    else:
        df_compteur = pd.DataFrame(columns=COLONNES_SUIVI_OFFICIELRES)

    if donnees_exclus_uniquement:
        df_exclus = pd.DataFrame(donnees_exclus_uniquement)
    else:
        df_exclus = pd.DataFrame(columns=COLONNES_SUIVI_OFFICIELRES)
    
    def styliser_tableau(val):
        color = 'background-color: #ffcccc; color: #cc0000; font-weight: bold;' if "EXCLU" in str(val) else ''
        return color
        
    st.dataframe(df_compteur.style.map(styliser_tableau, subset=["Situation Réglementaire"]), use_container_width=True)
    
    # --- SECTION EXPORTATION DES ETUDIANTS EXCLUS ---
    st.markdown("### 💾 Extraction et Téléchargement de la Liste des Exclus")
    
    if len(donnees_exclus_uniquement) == 0:
        st.info("Aucun étudiant n'a atteint le seuil d'exclusion (5 absences cumulées dans la même matière) pour le moment.")
    else:
        st.warning(f"Total détecté : {len(df_exclus)} cas d'exclusion réglementaire par matière.")
        
        col_btn_excel, col_btn_html = st.columns(2)
        
        with col_btn_excel:
            buffer_excel = io.BytesIO()
            with pd.ExcelWriter(buffer_excel, engine='openpyxl') as writer:
                df_exclus.to_excel(writer, index=False, sheet_name='Liste_des_Exclus_Matiere')
            buffer_excel.seek(0)
            
            st.download_button(
                label="📥 Télécharger la Liste des Exclus en format Excel (.xlsx)",
                data=buffer_excel,
                file_name=f"Liste_Exclus_Par_Matiere_{datetime.now().strftime('%d_%m_%Y')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            
        with col_btn_html:
            html_style = """
            <style>
                table { border-collapse: collapse; width: 100%; font-family: 'Calibri', sans-serif; margin-top: 20px; }
                th { background-color: #cc0000; color: white; padding: 10px; text-align: left; border: 1px solid #333; }
                td { padding: 8px; border: 1px solid #666; }
                tr:nth-child(even) { background-color: #f2f2f2; }
                .title { font-family: 'Arial', sans-serif; color: #333; font-size: 18px; font-weight: bold; text-align: center; margin-bottom: 5px; }
                .subtitle { font-family: 'Calibri', sans-serif; color: #555; text-align: center; font-size: 13px; margin-bottom: 25px; }
                .badge-exclu { color: #cc0000; font-weight: bold; background-color: #ffcccc; padding: 4px; border-radius: 3px; }
            </style>
            """
            html_titre_bloc = f'<div class="title">RÉPUBLIQUE ALGÉRIENNE DÉMOCRATIQUE ET POPULAIRE</div>'
            html_titre_bloc += f'<div class="subtitle">{TITRE_PLATEFORME}<br>LISTE DES EXCLUS UNITAIRES PAR COUPLE (ENSEIGNANT / MATIÈRE)</div>'
            
            html_table_brute = df_exclus.to_html(index=False)
            html_table_brute = html_table_brute.replace('❌ EXCLU DE LA MATIÈRE', '<span class="badge-exclu">❌ EXCLU DE LA MATIÈRE</span>')
            
            html_document_complet = f"<html><head><meta charset='utf-8'>{html_style}</head><body>{html_titre_bloc}{html_table_brute}</body></html>"
            
            st.download_button(
                label="🌐 Télécharger la Liste des Exclus en format HTML (.html)",
                data=html_document_complet,
                file_name=f"Liste_Exclus_Par_Matiere_{datetime.now().strftime('%d_%m_%Y')}.html",
                mime="text/html",
                use_container_width=True
            )

    # ==========================================
    # BOUTON STRICT DE PURGE DE L'HISTORIQUE
    # ==========================================
    st.divider()
    st.markdown("### 🗑️ Zone de Sécurité : Purge des Données Courantes")
    
    with st.expander("Ouvrir les options de nettoyage de la base de session"):
        st.warning("Attention : L'effacement de l'historique supprimera définitivement tous les enregistrements de justificatifs et réinitialisera les compteurs d'absences de tous les étudiants à zéro.")
        confirmation_purge = st.checkbox("Je confirme vouloir vider l'historique et remettre à zéro les compteurs d'exclusions.")
        
        if confirmation_purge:
            if st.button("⚠️ Exécuter l'effacement complet des données"):
                st.session_state["historique_justifications"] = []
                st.session_state["compteur_absences_strict"] = {}
                st.session_state["metadonnees_absences_strict"] = {}
                st.success("L'historique global et les compteurs d'absences ont été réinitialisés avec succès.")
                st.rerun()
